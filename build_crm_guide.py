# -*- coding: utf-8 -*-
"""Generates the detailed CRM User & Workflow Guide as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x1E, 0x40, 0xAF)
ACCENT = RGBColor(0x0F, 0x76, 0x6E)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RIPPLE = RGBColor(0x9A, 0x34, 0x12)   # burnt orange for ripple callouts

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)

def _cell(cell, text, bold=False, white=False, size=9.5, color=None):
    cell.text = ""
    p = cell.paragraphs[0]; r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif color: r.font.color.rgb = color

def _borders(tbl, hexc="D1D5DB", inside=True):
    edges = ("top","left","bottom","right") + (("insideH","insideV") if inside else ())
    tblPr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in edges:
        el = OxmlElement(f"w:{e}"); el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
        el.set(qn("w:space"),"0"); el.set(qn("w:color"),hexc); b.append(el)
    tblPr.append(b)

def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = BRAND
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    bo = OxmlElement("w:bottom"); bo.set(qn("w:val"),"single"); bo.set(qn("w:sz"),"6")
    bo.set(qn("w:space"),"4"); bo.set(qn("w:color"),"3B82F6"); pbdr.append(bo); pPr.append(pbdr)

def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = ACCENT

def h3(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = INK

def para(text, italic=False, size=10.5, color=None):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def lead(text, cont):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(cont); r2.font.size = Pt(10.5); return p

def bullet(text, cont=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    if cont:
        r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
        r2 = p.add_run(cont); r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text); r.font.size = Pt(10.5)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(10.5); return p

def box(label, text, fill, edge, lab_color):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; _shade(c, fill); c.paragraphs[0].text = ""
    r = c.paragraphs[0].add_run(label + "  "); r.bold = True; r.font.color.rgb = lab_color; r.font.size = Pt(10)
    r2 = c.paragraphs[0].add_run(text); r2.font.size = Pt(10); r2.font.color.rgb = INK
    _borders(tbl, edge, inside=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def ripple(text):
    box("RIPPLE EFFECT →", text, "FFF7ED", "FDBA74", RIPPLE)

def note(text):
    box("NOTE:", text, "EFF6FF", "BFDBFE", BRAND)

def example(text):
    box("WORKED EXAMPLE:", text, "F0FDF4", "BBF7D0", RGBColor(0x15,0x80,0x3D))

def table(headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, t in zip(tbl.rows[0].cells, headers):
        _shade(c, "1E40AF"); _cell(c, t, bold=True, white=True, size=9.5)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            _cell(cells[i], val, bold=(i == 0), size=9.5, color=INK if i == 0 else None)
    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = Inches(w)
    _borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ============================ COVER ============================
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SECURITY SERVICES CRM"); r.bold = True; r.font.size = Pt(32); r.font.color.rgb = BRAND
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("The Complete Workflow & Interaction Guide"); r.font.size = Pt(17); r.font.color.rgb = ACCENT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("What every action does — and exactly what it changes everywhere else"); r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = MUTED
for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A hands-on guide for the team who runs the system.\nEvery feature is explained as cause and effect: you do X here, and Y, Z and the bank balance change there.")
r.font.size = Pt(10.5); r.font.color.rgb = MUTED
doc.add_page_break()

# ============================ 1. BIG PICTURE ============================
h1("1. The Big Picture — Everything Is a Chain")
para("This CRM runs the whole back office of a guard-deployment / security-services business: the clients you guard, the contracts you sign, the guards you employ, their daily attendance, their monthly pay, the invoices you send, and every rupee that moves through your banks and cash box.")
para("The one idea that makes the entire system click: nothing stands alone. Every action quietly changes other screens. Learn the chain and the whole thing becomes predictable.")
h2("The two golden chains")
lead("Money OUT (cost side):  ", "CLIENT → CONTRACT → EMPLOYEE → ATTENDANCE → PAYROLL → BANK / CASH")
lead("Money IN (revenue side):  ", "CLIENT → CONTRACT → INVOICE → PAYMENT → BANK / CASH")
para("Everything else — Roster, Incidents, Inventory, Expenses, Partnership, Reports — either hangs off these two chains or reads them back to you.")
note("The single most important mechanic in the whole app: almost every money action writes a permanent line in a bank/cash ledger AND immediately moves the balance of the account it touched. Disburse a salary from a bank account and that account's balance drops by exactly that amount, that instant. Record a client payment into an account and it rises. This is covered in full in Section 4.")

# ============================ 2. ROLES ============================
h1("2. Who Can Do What — Roles & Permissions")
para("Two people can log into the same system and see totally different menus. That is deliberate — access is controlled per person.")
table(["Role", "What they are, and what they can reach"],
[
 ["Super Super Admin","The platform owner (us). Sits above every company; creates companies, runs their subscriptions, and can 'View as' any company to look inside. Not a day-to-day business user."],
 ["Super Admin","The owner / top manager of one company. Automatically has every permission. This is the account that runs the business."],
 ["HR","A staff login that only sees the permissions the Super Admin switches on — usually employees, attendance, payroll, documents."],
 ["Accounting","A staff login limited to the finance permissions granted — banks, expenses, invoices, reports, etc."],
], (1.4, 5.1))
h2("How permissions actually behave")
bullet("Each feature has a view switch (open & read) and an edit switch (add / change / delete). ")
bullet("If you lack a feature's permission, it does not appear in your sidebar at all — you cannot even reach it by typing the address. ", "A 'missing' screen is almost always a missing permission, not a bug.")
bullet("Super Admin and Super Super Admin bypass all switches — they always have everything.")
bullet("Branch scope: ", "a user can be tied to one branch, and then only sees that branch's data. The dashboard says 'scoped to your branch' when this applies.")
note("First login forces you to set your own password before anything else. You can change it later from the sidebar, and edit your name/photo from your profile at the bottom of the sidebar.")

# ============================ 3. DASHBOARD ============================
h1("3. Dashboard — The Cockpit (Read-Only)")
para("The first screen after login. It never accepts data entry; it only summarises what the other features already contain, and every tile links to the screen behind it.")
h3("What it surfaces")
bullet("Stat cards: active employees, today's attendance %, this-month expenses & payroll (vs last month), active contracts, open incidents, licences expiring within 30 days, roster gaps for the next 7 days.")
bullet("Live bank balances and their total; top 10 clients by payments received this month; 7-day attendance trend; expenses by category; and watch-lists for upcoming compliance, contracts ending, recent incidents, roster coverage and period-close status.")
ripple("Nothing here changes data — but every number is a live mirror. If a bank balance looks wrong on the dashboard, the cause is a transaction in Banks & Ledgers, not the dashboard. Fix causes at the source; the dashboard corrects itself. A Super Admin can hide any widget for the whole company from Settings.")

# ============================ 4. THE MONEY ENGINE ============================
h1("4. The Money Engine — How Balances Actually Move")
para("Read this section before any finance feature. Once you understand it, Payroll, Invoices, Expenses, Cheques and Cash Custody all behave the same predictable way.")
h2("Two pools of money")
bullet("Bank accounts — ", "each has its own running balance.")
bullet("Cash in hand (the treasury) — ", "one company cash balance, the physical cash box.")
h2("Every movement does two things at once")
para("Whenever money moves, the system (a) writes a permanent ledger line recording it, and (b) adjusts the balance of the account/cash it came from or went to. The two always happen together, so balances and history can never disagree.")
table(["Action", "Cash box", "Bank account", "Also happens"],
[
 ["Set opening balance on a bank","—","Set to opening amount","Opening ledger line written"],
 ["Cash deposit (cash → bank)","Goes DOWN","Goes UP","Two-sided ledger line"],
 ["Withdraw to cash (bank → cash)","Goes UP","Goes DOWN","Two-sided ledger line"],
 ["Wire transfer (bank → bank)","—","One down, one up","A paired transaction links both sides"],
 ["Disburse salary via Bank","—","Goes DOWN by net pay","'Payroll' ledger line on that account"],
 ["Disburse salary via Cash","Goes DOWN by net pay","—","'Payroll' (cash) ledger line"],
 ["Record client payment to Bank","—","Goes UP","'Receipt' line; invoice marked paid"],
 ["Pay a cash/bank expense","Down (if cash)","Down (if bank)","'Expense' ledger line"],
], (2.15, 1.05, 1.35, 1.65))
example("You disburse a guard's net pay of PKR 45,000 from 'HBL Main'. Instantly: HBL Main's balance drops by 45,000; a 'Payroll' ledger line is filed against HBL Main; the payslip is marked Disbursed & Cleared; the dashboard's bank total and this-month payroll both update. If you later un-disburse it, every one of those reverses exactly.")
note("Double-disbursement is blocked: the system 'claims' a payslip before moving money, so clicking twice cannot pay a guard twice. If any money step fails, the claim is released and nothing is left half-done.")
h2("Auto-zeroing accounts")
para("A bank account can be set to auto-zero monthly — handy for a pass-through account that should always start each month empty. At month rollover the system sweeps it back to zero automatically and records the adjustment.")

# ============================ 5. CLIENTS ============================
h1("5. Clients — The Master Anchor")
para("A Client is a company you guard. It is the anchor of the entire money-in side: 'every employee, contract and invoice anchors here.'")
h3("What you store")
bullet("Identity: auto client code, name, contacts, industry, branch.")
bullet("Tax profile: NTN, STRN, filer/non-filer, and a list of tax lines (e.g. sales tax added, income tax withheld) that will drive this client's invoices.")
bullet("Billing: type (Standard/SLA), invoice group (Fixed/Variable/SLA), billing address, authorised signatory + CNIC.")
bullet("Remittance accounts: the bank account(s) the client pays into; the default one prints on invoices.")
bullet("Employee ID prefix: e.g. 'ABC' → guards for this client are coded ABC-001, ABC-002…")
h3("What creating / editing a client unlocks")
ripple("Adding a client immediately makes it selectable everywhere downstream: you can now attach Contracts to it, assign Employees to it, raise Invoices for it, and it starts appearing in Receivables. Its tax profile pre-loads onto every invoice you raise for it. Its allowed-leaves and EOBI settings become the fallback the payroll engine uses when a contract doesn't specify them (see Section 10). Its ID prefix renames the employee codes of guards assigned to it — and every code change is kept in an ID-history trail on the employee.")

# ============================ 6. CONTRACTS ============================
h1("6. Contracts — The Deal, In Detail")
para("The client is 'who'; the contract is 'the deal'. One client can hold several contracts, each with its own headcount, rates, dates and terms.")
h3("What a contract holds")
bullet("Type: Guard Deployment (bill for people) or Services (bill for hardware — weapons/equipment).")
bullet("Dates: start, and either an end date or an 'open-ended' flag governed by a notice period.")
bullet("Contract lines — the core: ", "one line per category (Senior Supervisor, Supervisor, Guard, Reliever, Weapon, Equipment…), each with a committed headcount and a unit rate. Contract value = Σ (headcount × rate).")
bullet("Leave allowance & EOBI: set here per contract (blank = inherit the client's values).")
bullet("Status: Draft / Active / Expired / Terminated; plus annual escalation %, renewal terms, and the signed file (stored in Google Drive).")
h3("Addendums — changes over time")
para("Deals change: more guards, fewer guards, a new rate. Instead of editing the original, you file a dated addendum ('add 5 guards effective 1 March'). The system then computes the effective committed headcount on any date = original lines + all addendums up to that date — preserving a true history.")
ripple("Contract lines and addendums quietly power staffing checks and billing: (1) each guard can be assigned to a specific contract line, and the system compares committed headcount vs. actually-deployed guards per category, exposing over/under-staffing; (2) those same lines and rates become the line items when you generate an invoice; (3) the contract's leave-allowance and EOBI feed every assigned guard's payslip; (4) a contract nearing its end date automatically surfaces on the Dashboard and in Licences & Renewals.")

# ============================ 7. INVOICES ============================
h1("7. Invoices — Billing the Client")
para("An invoice is the formal request for payment. It can be built by hand or generated from a contract, and is tracked from raised to fully paid.")
h3("The workflow")
numbered("Create an invoice for a client. The system proposes the next number (e.g. INV-057) and rejects duplicates.")
numbered("It carries a subtotal, taxes (added and withheld, pulled from the client's tax profile), any previous outstanding balance carried forward, and a total due — with amount-in-words and the remit account, all printable.")
numbered("Status walks through Pending → Delivered → Unpaid → Partly-Paid → Paid as reality changes.")
numbered("Record payments against it; download a branded PDF on your company template; attach the signed copy (Drive).")
h3("What recording a payment does (the important part)")
para("Recording a payment runs as one atomic step — everything below commits together or not at all:")
bullet("The amount is allocated oldest-invoice-first across the client's unpaid invoices.")
bullet("Each affected invoice's amount-received rises and its status flips toward Paid.")
bullet("The chosen bank account (or cash) goes UP by the amount, with a 'Receipt' ledger line.")
bullet("The client's outstanding balance in Receivables drops accordingly.")
ripple("One client payment moves five things: the invoice status, the client's receivable balance, the receiving account's balance, the Cash Flow 'revenue' figure, and the Dashboard's 'Top clients' ranking — plus the Financial Reports revenue line. Auto-invoicing can also generate a batch of invoices from contracts in one run.")

# ============================ 8. EMPLOYEES ============================
h1("8. Employees — The Workforce Hub")
para("Every guard, supervisor and office staffer lives here. Attendance, payroll, roster, incidents, inventory and documents all point back to this record.")
h3("What a profile holds")
bullet("Identity & code: name, auto/client-prefixed code (with a full ID-history trail on reassignment), CNIC, DOB, father/husband name, blood group, addresses, emergency contact.")
bullet("Category: Client guard, Office staff, or Reliever (relievers get their own attendance & payroll screens).")
bullet("Assignment: which client, contract, and contract-line slot; the shift.")
bullet("Pay setup: base salary or per-day salary, plus a fixed allowance that is always paid, untaxed, regardless of attendance.")
bullet("Compliance: weapon & guard-service licence numbers/expiries, medical fitness expiry, EOBI reg, bank/IBAN, contract type (permanent/contract/probation/daily-wages) and probation end.")
bullet("Status: Active, On Leave, Inactive. An Inactive or On-Leave guard frees up their contract-line slot.")
ripple("Changing an employee ripples widely: setting their status to Inactive frees their contract-line slot (changing staffing counts on the Contract) and drops them from active-headcount stats; changing their client re-codes them under the new client's prefix and logs it in ID history; their licence expiries feed the Licences watch-list and the Dashboard's 'expiring <30d' card; their salary fields drive next month's payslip.")

# ============================ 9. ATTENDANCE ============================
h1("9. Attendance — The Daily Heartbeat")
para("Attendance is the day-by-day record of who showed up. Small entry, huge downstream effect — it is the single biggest driver of each guard's pay.")
h3("How it works")
bullet("Each employee is marked Present, Absent or Leave per day. Mark a whole team fast, or open one guard's monthly calendar and bulk-mark.")
bullet("Extra detail: half-day, late arrival, hours worked, overtime.")
bullet("Relievers are special: ", "because a reliever can cover a different client each day, each reliever day records which client it was worked for. That per-day attribution is what later splits a reliever's pay across the clients they covered.")
ripple("Every mark you make here is money next month: present/absent/leave counts flow straight into the payslip. Leave within the contract/client allowance is paid; leave beyond it is treated as unpaid absence and docks pay. The same marks drive the Dashboard's attendance % and 7-day trend. The Roster is the plan; Attendance is what actually happened — the two should match.")

# ============================ 10. PAYROLL ============================
h1("10. Payroll — Attendance Becomes Money")
para("Pick a month, the system builds one payslip per employee, you review and adjust, then disburse.")
h3("How each payslip is built")
bullet("Working / present / absent / leave days come from that month's attendance.")
bullet("Allowed leaves (from the Contract, else the Client) are applied: leave within allowance is paid; leave beyond it is docked like absence.")
bullet("Base (or per-day) salary is prorated by attendance; the fixed allowance is added on top untaxed; then bonus, deductions, income tax, EOBI and any outstanding advance are applied to reach net salary.")
bullet("EOBI amount also comes from the contract/client settings.")
h3("Disbursement — where the money actually moves")
bullet("Each payslip is paid via Cash, Bank or Cheque.")
bullet("Bank: ", "the chosen account's balance drops by the net pay and a 'Payroll' ledger line is filed. Un-disbursing reverses it exactly.")
bullet("Cash: ", "the cash box drops by the net pay, with a cash 'Payroll' line.")
bullet("Cheque: ", "the bank balance is NOT touched at disbursement — the linked cheque clearing handles the bank side later (see Section 13).")
bullet("Relievers have their own payroll screen; pay is attributed across the clients each reliever covered that month.")
ripple("Payroll pulls from three features and pushes to two: it reads Attendance (days), the Contract/Client (leave & EOBI), and Expenses (an outstanding advance is auto-deducted so you never double-pay). It then pushes into Banks/Cash (balance down) and into Cash Flow & Reports (payroll outflow). Disbursed payslips can be exported to PDF.")
example("A guard has base 40,000, allowance 5,000, took 1 leave over a 2-leave allowance (so within allowance — fully paid), and has a 3,000 advance outstanding. Net = 40,000 + 5,000 − EOBI − 3,000 (advance). You pay via Bank 'UBL Payroll' → UBL Payroll drops by the net, the advance in Expenses is cleared, and the payslip shows Disbursed.")

# ============================ 11. RELIEVERS ============================
h1("11. Relievers — Floating Guards, Attributed Fairly")
para("Relievers fill gaps across many clients rather than sitting at one post, so they get dedicated Attendance and Payroll screens under the Relievers group. Mechanics are identical to normal attendance and payroll, with one twist: every reliever day is tagged with the client it covered, so their monthly pay is split and attributed to each client's cost — keeping per-client profitability honest.")

# ============================ 12. BANKS & LEDGERS ============================
h1("12. Banks & Ledgers — The Money Hub (4 Tabs)")
para("The financial heart of the CRM. One screen, four tabs; nearly every rupee lands here.")
h2("Tab 1 — Bank Accounts")
bullet("Every account with live balance, opening balance, IBAN, type; owned by the company, a partner, or a client.")
bullet("Actions: cash deposit (cash→bank), withdraw to cash (bank→cash), wire transfer (bank→bank, recorded as a linked pair), edit opening balance, per-account transaction log, exportable statements and deposit slips.")
bullet("Each action moves balances exactly as described in Section 4 and files a ledger line.")
h2("Tab 2 — Client Receivables")
para("A live list of what each client still owes — built from unpaid & partly-paid invoices plus opening balances. Record a payment here or view/export a full client statement. This is the money-in ledger.")
h2("Tab 3 — Accounts Payable")
para("What your company owes. Expenses entered as 'Payable' sit here until you mark them paid (via cash or bank) — at which point the money leaves the chosen account and a ledger line is filed. This is the pay-it-later ledger.")
h2("Tab 4 — Cash Custody")
para("Tracks physical cash held by the company and by cash handlers. Deposits/withdrawals on the Bank Accounts tab sync here automatically. Covered in full in Section 20.")
ripple("Banks & Ledgers is the destination of both golden chains. Invoice payments land in Receivables and raise an account; payroll and paid expenses draw accounts down; deposits/withdrawals/transfers shuffle money between cash and banks. Everything it records flows onward into Cash Flow, Financial Reports and the Chart of Accounts' journal.")

# ============================ 13. CHEQUES ============================
h1("13. Cheques — How They Clear")
para("Cheques get special handling because a written cheque isn't cash in the account yet — it clears later. The system models exactly that.")
h3("What a cheque is")
bullet("It has a number, amount, date, a direction (Outgoing = you pay out, Incoming = you receive) and a type, plus a status of Pending or Cleared, and an optional scanned attachment (Drive).")
h3("Cheques link to real payments")
para("A cheque can be attached to a payslip, an expense, an advance, or an invoice payment. The system continuously adds up everything linked to each cheque, so you can see how much of the cheque's value has been used and whether it is fully accounted for.")
h3("The clearing rule (why balances stay right)")
bullet("Outgoing / payment cheques: ", "when you disburse a salary or pay an expense 'by Cheque', the bank balance is NOT reduced yet — the cheque is a pending promise. The account is only reduced when the cheque clears. This is why a cheque-paid payslip doesn't move the bank at disbursement time.")
bullet("Incoming / deposit cheques: ", "a cheque you receive doesn't raise the bank balance until it clears either.")
ripple("Cheques decouple 'I promised to pay' from 'the bank actually moved'. Until a cheque clears, your bank balance reflects reality (money still in the account) while the payslip/expense already shows as handled. You filter cheques by Pending vs Cleared to see outstanding promises. Pending outgoing cheques are effectively money spoken-for but not yet gone.")

# ============================ 14. EXPENSES ============================
h1("14. Expenses — Operating Costs & Advances")
para("What the business spends that isn't payroll — weapons, uniforms, fuel, rent, utilities, insurance, licences, EOBI/PESSI/IESSI, taxes — plus employee advances.")
h3("How an expense behaves by payment mode")
bullet("Cash or Bank: ", "pays immediately — the cash box or chosen account drops now, with an 'Expense' ledger line.")
bullet("Payable: ", "parks in Accounts Payable (Section 12) and moves no money until you mark it paid.")
bullet("Cheque: ", "links a cheque; the bank moves only when that cheque clears (Section 13).")
bullet("Each expense carries a category, a P&L bucket (cost of services vs operating expense), optional client/branch/vendor, and a Drive receipt. Vendors and custom categories are managed inline.")
h3("Advances")
para("An advance is money fronted to an employee, recorded with a payment mode and account (so it also draws that account down now).")
ripple("Expenses touch three other features: (1) a Payable sits in Banks & Ledgers until settled; (2) a cash/bank expense immediately lowers that account's balance; (3) an outstanding advance is automatically netted off that employee's next payslip in Payroll — so the money you fronted is recovered without any manual deduction. Categorised expenses also drive the Dashboard expense pie and the P&L cost lines.")

# ============================ 15. CASH FLOW ============================
h1("15. Cash Flow — Money In vs Money Out")
para("A read-only netting view for a chosen month, range, or all time: payments received, disbursed net salaries, cash/bank expenses plus paid payables, and advances — arriving at Revenue − Payroll − Expenses − Advances. It creates nothing; it re-reads the ledgers the features above produced.")

# ============================ 16. FINANCIAL REPORTS ============================
h1("16. Financial Reports — P&L, Partnership & Statements")
para("The reporting centre: it reads revenue (invoice payments), payroll and expenses back as a Profit & Loss, shows the partnership position, and prints full client statements. Reports never create data — their accuracy depends entirely on everything upstream being entered honestly and on time.")

# ============================ 17. CHART OF ACCOUNTS ============================
h1("17. Chart of Accounts — The Accountant's View")
para("For finance users who think in double-entry. It lists every account (assets, liabilities, equity, revenue, expenses), shows a Trial Balance built from the system's double-entry journal, and lets you drill into any account's General Ledger; you can also post manual journal entries. Most users never open this — but every invoice payment, payroll run and expense quietly writes journal lines that appear here, which is what makes the Trial Balance self-maintaining.")

# ============================ 18. PERIOD CLOSE ============================
h1("18. Period Close — Locking the Books")
para("Once a month is final, you close it. A closed month is locked: no new edits, payments or entries can land in it, which protects reported figures from accidental after-the-fact changes. Re-opening requires an explicit confirmation and is recorded. The Dashboard shows whether the current month is open or closed.")
ripple("Closing a period is a guard rail across the whole app: it blocks writes into that month everywhere — Payroll, Invoices, Expenses, Banks — so a report you signed off can't silently change underneath you.")

# ============================ 19. PARTNERSHIP FINANCE (FULL) ============================
h1("19. Partnership Finance — Partners, Profit Split & Projects")
para("If the business is owned by partners (and possibly outside investors), this section tracks every rupee of their stake. It has three linked features. Read them in order — they build on each other.")

h2("19a. Partner Accounts — the running ledger per partner")
para("Each partner has a personal running ledger with the company. Its balance answers one question: does the company owe the partner money, or has the partner taken out more than they're owed?")
h3("The four entry types and what each does to the balance")
table(["Entry", "Meaning", "Effect on partner balance"],
[
 ["Opening Balance","Where the partner started (locks once set)","Sets the starting point"],
 ["Profit Allocation","The partner's share of profit for a period (can be negative for a loss share)","INCREASES (company owes them more)"],
 ["Contribution","Partner puts their own money INTO the company","INCREASES (company owes them more)"],
 ["Drawing","Partner takes money OUT of the company","DECREASES (reduces what they're owed)"],
], (1.6, 2.9, 2.1))
para("Balance = Opening + Allocations + Contributions − Drawings. A positive balance means the company owes the partner; a negative balance means the partner is overdrawn (has taken more than their share).")
bullet("Scope: ", "a partner is either Company-wide (an owner) or tied to a Branch (a regional partner who only shares that branch's results).")
bullet("Allocation method: ", "Fixed % (their share auto-computed from a set percentage) or Manual (you type each period's share).")
bullet("Drawings and Contributions do NOT touch the Profit & Loss ", "— they are movements of ownership money, not business costs or income. Only Profit Allocations reflect actual profit.")
example("A partner opens at 0, is allocated 200,000 profit for the month, then draws 120,000 in cash. Their ledger reads: 0 + 200,000 − 120,000 = 80,000. The company still owes this partner 80,000. Each line shows a running balance and the payment method used for drawings/contributions; the whole statement exports to CSV.")

h2("19b. Profit Distribution — the RULES that decide each partner's share")
para("Partner Accounts records the split after it's decided. Profit Distribution is where you define HOW profit is split in the first place, so allocations are consistent and defensible.")
h3("Distribution rules")
bullet("A rule applies at a level: ", "Company (all branches), a specific Branch, or a specific Client.")
bullet("Each rule has effective-from date and a set of lines: ", "each line sends a percentage to either a Partner or to 'Retained' (kept in the business). The lines must total no more than 100%.")
bullet("The most specific rule wins: ", "a Client-level rule overrides a Branch rule, which overrides the Company rule — so you can bill most work one way but carve out special clients.")
h3("Referral arrangements")
para("A partner who brings in business can earn a referral cut. You record: the referring partner, the source branch (and optionally a specific client), a percentage, and how it's funded:")
bullet("Off the top: ", "taken from profit before the normal split.")
bullet("Partners only: ", "funded from the partners' shares, not retained earnings.")
bullet("Custom split: ", "a bespoke arrangement.")
ripple("Profit Distribution is the policy layer; Partner Accounts is the ledger it feeds. You define the rules and referrals here, then post the resulting amounts as Profit Allocations in each partner's account (19a). The profit being divided comes from Financial Reports' P&L (Section 16) — which itself comes from all the revenue and cost features. So a single guard's attendance eventually influences what a partner is allocated three layers up.")

h2("19c. Project Financing — outside investors in specific projects")
para("Separate from the core partnership, some projects are funded by investors (who may be partners or third parties) expecting a return. This feature tracks that capital and what's owed back.")
h3("The pieces")
bullet("Project: ", "what's being funded — with a total required, a reserved-profit %, a payout gate (pay returns from company cash flow, or only from the project's own cash flow), and a status (Raising → Active → Completed).")
bullet("Investor: ", "a Partner or a Third Party; a partner-investor can be linked to their partner record.")
bullet("Investment: ", "an investor's stake in a project, either Profit Share (they get a cut of profit) or Fixed Finance (they get a fixed cost/return).")
h3("The investor ledger — six entry types")
table(["Entry", "What it records"],
[
 ["Capital In","Investor puts money into the project"],
 ["Capital Repayment","Their capital is returned"],
 ["Return Allocation","Their profit share is credited (accrued)"],
 ["Return Payout","That return is actually paid out"],
 ["Finance Cost Accrual","A fixed-finance charge is accrued"],
 ["Finance Cost Payment","That charge is paid"],
], (2.3, 4.3))
ripple("Project Financing keeps investor money ring-fenced from partner money. The outstanding capital and returns owed to investors are counted as company liabilities — and appear in the Cash Custody reconciliation (Section 20) so your cash position accounts for money that ultimately belongs to investors. The 'payout gate' decides whether returns can be paid from general company cash or only once the project itself generates cash.")

# ============================ 20. CASH CUSTODY ============================
h1("20. Cash Custody — Who Physically Holds the Cash")
para("Lives as the fourth tab of Banks & Ledgers. Where Bank Accounts track money in banks, Cash Custody tracks physical cash and who is holding it.")
h3("The pieces")
bullet("Cash locations: ", "a petty-cash float, or a custodian (a partner or staff member) who holds company cash, each with an opening balance; a location can also mirror a bank account.")
bullet("Custody transfers: ", "move cash from one location/handler to another, with a date and note.")
h3("The Cash Position & Reconciliation")
para("The Position tab lays out every place company money sits — 'Cash in Hand' (the treasury), each custody location, each partner's balance, and investor liabilities — so you can reconcile: does the cash we can see equal the cash we should have, once partner balances and investor money are accounted for?")
ripple("Cash Custody ties the physical world back to the ledgers: the company treasury figure it shows is the same cash box that Payroll (cash disbursements) and Expenses (cash payments) draw down; bank-side deposits/withdrawals sync in automatically; partner balances come from Partner Accounts (19a); investor liabilities come from Project Financing (19c). It is the one screen that answers 'where is all our cash, and does it add up?'")

# ============================ 21. ROSTER ============================
h1("21. Deployment Roster — Planning Who Goes Where")
para("The operational planner: a grid of employees against upcoming days, where each cell assigns a guard to a post and shift. 'Click any cell to assign.'")
bullet("Posts (deployment sites) belong to a client and contract and state how many guards they need and on what shift.")
bullet("Each cell has a status: Assigned, Confirmed, Leave Requested, Reliever Needed, or Unassigned — so gaps are visible at a glance.")
ripple("The roster reads Employees and Posts (from Contracts) and pushes coverage figures to the Dashboard ('roster gaps / next 7 days' and 'coverage %'). A 'Reliever Needed' cell is the signal that hands off to the Reliever workflow. What you plan here should later match what Attendance records as reality — divergence is the early warning of a no-show (which may become an Incident).")

# ============================ 22. INCIDENTS ============================
h1("22. Incidents — The Security Event Log")
para("A structured record of things that go wrong on site: theft, altercation, guard injury, weapon discharge, no-show, asset damage, client complaint, or other. Each gets a code, severity (low → critical), date/time, and a status (Open → Under Investigation → Resolved → Closed).")
ripple("An incident links to the Client and Post where it happened AND names the specific guards involved (from Employees) — giving you a per-guard and per-client incident history. You record whether/when the client was notified and the action taken; supporting files attach via Drive. Open and recent incidents surface on the Dashboard, and a spike of incidents against one client or guard is exactly the signal management needs.")

# ============================ 23. INVENTORY ============================
h1("23. Inventory — Weapons & Uniforms")
para("Tracks physical assets — weapons and uniforms — plus who holds them. Each item has a type, serial number, quantity, location/branch, licence expiry (weapons), and a status of Available / Issued / Maintenance.")
ripple("Issuing an item ties it to an Employee with an issue date; returning it records a condition (Good/Fair/Damaged) and frees it. Weapon licence expiries feed the Licences & Renewals watch-list. Buying weapons or uniforms shows up as an Expense — so inventory, employees, compliance and expenses all meet here.")

# ============================ 24. LICENCES ============================
h1("24. Licences & Renewals — Nothing Expires Unnoticed")
para("One sorted list of every expiring item across the whole business — guard weapon licences, guard-service licences, medical-fitness certificates, probation ends, contract end dates and company compliance items — ordered by days remaining.")
ripple("This screen creates nothing of its own; it aggregates expiry dates from Employees, Contracts, Inventory and the Compliance Calendar into one countdown. It is the safety net that stops a lapsed weapon licence or an expired contract slipping through — the same expiries feed the Dashboard's 'expiring <30d' card.")

# ============================ 25. COMPLIANCE ============================
h1("25. Compliance Calendar — Important Dates & Reminders")
para("Where you record important dates (licence renewals, tax filings, HR/payroll deadlines), each with a category, priority, and how many days ahead to be warned. You can also set recurring reminders (daily/weekly/monthly/yearly). These items — plus contract endings — are what power the Dashboard's 'Upcoming Compliance' alerts and feed the Licences list.")

# ============================ 26. DOCUMENTS ============================
h1("26. Documents — The Employee File Cabinet")
para("A per-employee document repository backed by Google Drive: upload contracts, CNIC copies, certificates and any other files against a specific employee, then download or delete them later. It's the same Drive storage the attachment buttons elsewhere (invoices, expenses, incidents, contracts, cheques) use behind the scenes.")

# ============================ 27. TASKS ============================
h1("27. Tasks — The Team To-Do Board")
para("A simple Kanban board (To Do / In Progress / Done) for assigning work and tracking progress. Admins assign and track; staff see and update what's theirs. It's the one feature open to everyone regardless of other permissions, so the whole team can coordinate.")

# ============================ 28. USERS ============================
h1("28. Users & Permissions — Access Control")
para("Where the Super Admin creates staff accounts and decides exactly what each can see and do.")
bullet("Create an HR or Accounting user; set name, title, branch and a temporary password.")
bullet("Tick the view/edit permission switches per feature — these are precisely what make sidebar items appear or disappear for that user.")
bullet("Scope a user to a branch so they only see that branch's data.")
ripple("This is the control room behind Section 2. Every 'why can't I see X?' question is answered here. Granting a permission makes a whole feature appear in that user's sidebar on their next load; revoking it makes the feature — and the ability to reach it — vanish.")

# ============================ 29. AUDIT LOG ============================
h1("29. Audit Log — The Complete Paper Trail")
para("Every change to every important record — employees, clients, contracts, invoices, payments, expenses, payslips, advances, cheques, bank accounts and transactions, branches, users, accounts, periods, posts, incidents and roster — is captured automatically: who changed it, when, and the exact before/after values. Only Super Admins and the platform owner can view it. Nothing needs switching on; it records silently, which makes it the final answer to 'who changed this number, and what was it before?'")

# ============================ 30. SETTINGS ============================
h1("30. Settings — Company Configuration")
para("Company-level setup: manage locations/regions and branches; choose which Dashboard widgets everyone sees; and set the company profile, invoice-template fields (what prints on invoice PDFs), brand theme/colour, logo, legal address, tax NTN, presentation currency and fiscal year.")
ripple("Settings quietly shapes many screens: branches defined here are the branch options across every feature; the invoice template here controls what your invoice PDFs show; hidden widgets here change everyone's Dashboard.")

# ============================ 31. AI + SSA ============================
h1("31. AI Assistant & Platform Owner View")
h2("The AI Assistant")
para("A floating chat button at the bottom-right of every screen. Ask about your own data in plain language ('how many guards are active for client X?', 'what did we spend on fuel last month?') and it answers by reading — never changing — your records. It respects your permissions and can only ever see your own company's data.")
h2("Platform Owner View (Super Super Admin only)")
para("Lists every company with employee/user counts; create companies, activate/deactivate them, and manage each company's subscription (record payments, extend the expiry by a number of days). An expired subscription automatically deactivates a company's access. The owner can 'View as' any company — dropping into its panel with a clear 'you are viewing X' banner — then exit back out.")

# ============================ 32. MASTER MAP ============================
h1("32. Master Map — Every Action and Its Ripple")
para("Read each row as: do the left, and these things change on their own elsewhere.")
table(["When you…", "…this is what changes elsewhere"],
[
 ["Add a Client","Unlocks its contracts, invoices and guard assignments; its tax profile pre-loads onto its invoices; its leave/EOBI become the payroll fallback."],
 ["Add a Contract","Sets committed headcount, rates, leave & EOBI; enables invoice generation, roster posts, and staffing-gap checks for the client."],
 ["File a Contract Addendum","Changes effective headcount/rate from a date forward without losing history; re-checks staffing counts."],
 ["Add / deactivate an Employee","Fills or frees a contract-line slot (changing staffing counts); re-codes on client change; feeds licence & headcount stats."],
 ["Mark Attendance","Sets next month's payslip (present/absent/leave, leave-over-allowance docks pay); drives dashboard attendance."],
 ["Disburse Payroll (Bank/Cash)","That account/cash drops by net pay; a payroll ledger line is filed; any advance is netted off; Cash Flow & Reports update."],
 ["Disburse Payroll (Cheque)","No bank move yet — the linked cheque clears the bank side later."],
 ["Record an Advance","Draws the paying account down now; auto-deducts from the employee's next payslip."],
 ["Raise an Invoice","Bills the client from contract lines + tax profile; adds to Receivables."],
 ["Record an Invoice Payment","Allocates oldest-first, marks invoices paid, raises the receiving account, cuts the client's receivable — all atomically; updates Top Clients & revenue."],
 ["Enter a cash/bank Expense","Draws that account down now; feeds the expense pie and P&L cost lines."],
 ["Enter a Payable Expense","Parks in Accounts Payable; money moves only when you mark it paid."],
 ["Write a Cheque","Nothing leaves the bank until the cheque clears; tracked as a pending promise."],
 ["Deposit / Withdraw / Transfer","Moves money between cash and banks (transfers as a linked pair); cash side syncs to Cash Custody."],
 ["Post a Profit Allocation / Drawing / Contribution","Moves a partner's running balance; allocations reflect profit, drawings/contributions don't hit P&L."],
 ["Set a Distribution Rule / Referral","Defines how future profit splits among partners/retained; the amounts get posted as allocations."],
 ["Record Project Investor Capital / Returns","Tracks investor liabilities; surfaces in the Cash Custody reconciliation."],
 ["Log an Incident","Links client, post and named guards; records client notification; hits the Dashboard."],
 ["Plan the Roster","Assigns guards to posts by shift; exposes gaps & 'reliever needed'; feeds coverage stats."],
 ["Issue Inventory","Ties an asset to an employee; weapon licences feed the Licences watch-list."],
 ["Close a Period","Locks the month app-wide so no feature can alter its numbers."],
 ["Everything above","Is read back — unchanged — by Cash Flow, Financial Reports, Chart of Accounts, Partnership Finance, Cash Custody, the Dashboard, the Audit Log and the AI Assistant."],
], (2.15, 4.35))

h2("The one rule to remember")
box("GARBAGE IN, GARBAGE EVERYWHERE.",
    "Because every screen feeds the next, a wrong attendance mark becomes a wrong payslip becomes a wrong bank balance becomes a wrong report becomes a wrong partner allocation. Enter data carefully and promptly at the source, and the whole system stays honest on its own.",
    "FEF2F2", "FECACA", RGBColor(0xB9,0x1C,0x1C))

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— End of Guide —"); r.italic = True; r.font.color.rgb = MUTED; r.font.size = Pt(10.5)

out = r"c:\Users\Abuzar\Desktop\employee-manager\CRM_User_and_Workflow_Guide.docx"
doc.save(out)
print("SAVED:", out)
